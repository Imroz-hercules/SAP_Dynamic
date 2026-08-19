"""
Generate KPI Formulas Word Document
Creates a comprehensive Word document with all KPI formulas, validation code, calculation code, and examples.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Add title
title = doc.add_heading('KPI Formulas Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Complete Reference with Code Examples')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ============================================================================
# MILLING KPIs
# ============================================================================

doc.add_heading('MILLING KPIs', 1)

# KPI 1: Mill Throughput
doc.add_heading('1. Mill Throughput (%)', 2)
doc.add_paragraph('Formula: (cap_per_h / nameplate_tph) × 100')
doc.add_paragraph('Scales Used: WG202, run_hours')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
code.style = 'Normal'
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 231-241\n'
    'nameplate_tph = 25.0\n'
    '\n'
    '# Calculate capacity per hour\n'
    'if cap_per_h <= 0 and run_hours > 0:\n'
    '    cap_per_h = WG202 / run_hours if WG202 > 0 else 0.0\n'
    '\n'
    '# Mill Throughput (%)\n'
    'if cap_per_h > 0:\n'
    '    mill_throughput = (cap_per_h / nameplate_tph * 100.0)\n'
    '    # Cap at reasonable maximum (150% of nameplate)\n'
    '    mill_throughput = min(mill_throughput, 150.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG202 = 200 tons, run_hours = 8 hours', style='List Bullet')
doc.add_paragraph('Step 1: Validation - has_milling_data = True', style='List Bullet')
doc.add_paragraph('Step 2: Calculate cap_per_h = 200 / 8 = 25 t/h', style='List Bullet')
doc.add_paragraph('Step 3: mill_throughput = (25 / 25) × 100 = 100%', style='List Bullet')
doc.add_paragraph('Step 4: Apply cap - min(100, 150) = 100%', style='List Bullet')
doc.add_paragraph('Result: 100%', style='List Bullet')

# KPI 2: Mill Time Efficiency
doc.add_heading('2. Mill Time Efficiency (%)', 2)
doc.add_paragraph('Formula: (run_hours / daily_hrs) × 100')
doc.add_paragraph('Scales Used: run_hours, daily_hrs')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 244-247\n'
    'if daily_hrs > 0:\n'
    '    mill_time_eff = (run_hours / daily_hrs * 100.0)\n'
    '    # Cap at 100% maximum\n'
    '    mill_time_eff = min(mill_time_eff, 100.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: run_hours = 8 hours, daily_hrs = 24 hours', style='List Bullet')
doc.add_paragraph('Step 1: Validation - has_milling_data = True', style='List Bullet')
doc.add_paragraph('Step 2: mill_time_eff = (8 / 24) × 100 = 33.33%', style='List Bullet')
doc.add_paragraph('Step 3: Apply cap - min(33.33, 100) = 33.33%', style='List Bullet')
doc.add_paragraph('Result: 33.33%', style='List Bullet')

# KPI 3: Total Utilization
doc.add_heading('3. Total Utilization (%)', 2)
doc.add_paragraph('Formula: (mill_time_eff × mill_throughput) / 100')
doc.add_paragraph('Scales Used: mill_time_eff, mill_throughput (calculated)')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 250-252\n'
    'total_util = (mill_time_eff * mill_throughput) / 100.0\n'
    '# Cap at reasonable maximum\n'
    'total_util = min(total_util, 150.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: mill_time_eff = 33.33%, mill_throughput = 100%', style='List Bullet')
doc.add_paragraph('Step 1: total_util = (33.33 × 100) / 100 = 33.33%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(33.33, 150) = 33.33%', style='List Bullet')
doc.add_paragraph('Result: 33.33%', style='List Bullet')

# KPI 4: Milling Gain
doc.add_heading('4. Milling Gain (%)', 2)
doc.add_paragraph('Formula: ((WG501 + WG502 + WG503 + WG301 + WG302) / WG201) × 100')
doc.add_paragraph('Scales Used: WG201, WG501, WG502, WG503, WG301, WG302')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 255-260\n'
    'if WG201 > 0:\n'
    '    total_output = WG501 + WG502 + WG503 + WG301 + WG302\n'
    '    if total_output > 0:\n'
    '        milling_gain = (total_output / WG201 * 100.0)\n'
    '        # Cap at reasonable maximum (120%)\n'
    '        milling_gain = min(milling_gain, 120.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG201 = 180t, WG501 = 120t, WG502 = 30t, WG503 = 48t, WG301 = 2t, WG302 = 1.5t', style='List Bullet')
doc.add_paragraph('Step 1: total_output = 120 + 30 + 48 + 2 + 1.5 = 201.5 tons', style='List Bullet')
doc.add_paragraph('Step 2: milling_gain = (201.5 / 180) × 100 = 111.94%', style='List Bullet')
doc.add_paragraph('Step 3: Apply cap - min(111.94, 120) = 111.94%', style='List Bullet')
doc.add_paragraph('Result: 111.94%', style='List Bullet')

# KPI 5: Milling Screening
doc.add_heading('5. Milling Screening (%)', 2)
doc.add_paragraph('Formula: (WG301 / WG201) × 100')
doc.add_paragraph('Scales Used: WG301, WG201')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 264-267\n'
    'if WG201 > 0 and WG301 > 0:\n'
    '    screening_ratio = (WG301 / WG201 * 100.0)\n'
    '    # Cap at reasonable maximum (20%)\n'
    '    screening_ratio = min(screening_ratio, 20.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG301 = 2 tons, WG201 = 180 tons', style='List Bullet')
doc.add_paragraph('Step 1: screening_ratio = (2 / 180) × 100 = 1.11%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(1.11, 20) = 1.11%', style='List Bullet')
doc.add_paragraph('Result: 1.11%', style='List Bullet')

# KPI 6: Flour Extraction
doc.add_heading('6. Flour Extraction (%)', 2)
doc.add_paragraph('Formula: ((WG501 + WG502) / WG202) × 100')
doc.add_paragraph('Scales Used: WG501, WG502, WG202')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 272-277\n'
    'if WG202 > 0 and (WG501 > 0 or WG502 > 0):\n'
    '    total_flour = WG501 + WG502\n'
    '    if total_flour > 0:\n'
    '        flour_extraction = (total_flour / WG202 * 100.0)\n'
    '        # Cap at reasonable maximum (85%)\n'
    '        flour_extraction = min(flour_extraction, 85.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG501 = 120t, WG502 = 30t, WG202 = 200t', style='List Bullet')
doc.add_paragraph('Step 1: total_flour = 120 + 30 = 150 tons', style='List Bullet')
doc.add_paragraph('Step 2: flour_extraction = (150 / 200) × 100 = 75%', style='List Bullet')
doc.add_paragraph('Step 3: Apply cap - min(75, 85) = 75%', style='List Bullet')
doc.add_paragraph('Result: 75%', style='List Bullet')

# KPI 7: Milling Loss
doc.add_heading('7. Milling Loss (%)', 2)
doc.add_paragraph('Formula: ((WG202 - (WG501 + WG502 + WG503)) / WG202) × 100')
doc.add_paragraph('Scales Used: WG202, WG501, WG502, WG503')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 285-290\n'
    'if WG202 > 0:\n'
    '    total_output = WG501 + WG502 + WG503\n'
    '    if total_output > 0:\n'
    '        milling_loss = ((WG202 - total_output) / WG202 * 100.0)\n'
    '        # Ensure it\'s not negative\n'
    '        milling_loss = max(milling_loss, 0.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG202 = 200t, WG501 = 120t, WG502 = 30t, WG503 = 48t', style='List Bullet')
doc.add_paragraph('Step 1: total_output = 120 + 30 + 48 = 198 tons', style='List Bullet')
doc.add_paragraph('Step 2: milling_loss = ((200 - 198) / 200) × 100 = 1%', style='List Bullet')
doc.add_paragraph('Step 3: Ensure non-negative - max(1, 0) = 1%', style='List Bullet')
doc.add_paragraph('Result: 1%', style='List Bullet')

# KPI 8: Net Hours
doc.add_heading('8. Net Hours (hrs)', 2)
doc.add_paragraph('Formula: max(run_hours - downtime, 0.0)')
doc.add_paragraph('Scales Used: run_hours, downtime')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Line 204\n'
    'net_hours = max(run_hours - downtime, 0.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: run_hours = 8 hours, downtime = 0.5 hours', style='List Bullet')
doc.add_paragraph('Step 1: net_hours = 8 - 0.5 = 7.5 hours', style='List Bullet')
doc.add_paragraph('Step 2: Ensure non-negative - max(7.5, 0) = 7.5 hours', style='List Bullet')
doc.add_paragraph('Result: 7.5 hours', style='List Bullet')

# KPI 9: Downtime
doc.add_heading('9. Downtime (hrs)', 2)
doc.add_paragraph('Formula: Direct SCADA value')
doc.add_paragraph('Scales Used: downtime')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# No validation - direct value from SCADA'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Line 133\n'
    'downtime = safe(row.get("WG202_Stop_Start"), 0.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: downtime = 0.5', style='List Bullet')
doc.add_paragraph('Result: 0.5 hours', style='List Bullet')

# KPI 10: Max Utilization of Milling Capacity
doc.add_heading('10. Max Utilization of Milling Capacity (%)', 2)
doc.add_paragraph('Formula: (WG202 / (run_hours × 25.0)) × 100')
doc.add_paragraph('Scales Used: WG202, run_hours')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 296-299\n'
    'if run_hours > 0:\n'
    '    max_utilization_milling_capacity = (WG202 / (run_hours * 25.0)) * 100.0\n'
    '    # Cap at reasonable maximum (150%)\n'
    '    max_utilization_milling_capacity = min(max_utilization_milling_capacity, 150.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG202 = 200 tons, run_hours = 8 hours', style='List Bullet')
doc.add_paragraph('Step 1: max_utilization = (200 / (8 × 25)) × 100 = 100%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(100, 150) = 100%', style='List Bullet')
doc.add_paragraph('Result: 100%', style='List Bullet')

# KPI 11: Pre Cleaning Screening
doc.add_heading('11. Pre Cleaning Screening (%)', 2)
doc.add_paragraph('Formula: (WG302 / WG101) × 100')
doc.add_paragraph('Scales Used: WG302, WG101')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 303-306\n'
    'if WG101 > 0 and WG302 > 0:\n'
    '    pre_cleaning_screening = (WG302 / WG101) * 100.0\n'
    '    # Cap at reasonable maximum (20%)\n'
    '    pre_cleaning_screening = min(pre_cleaning_screening, 20.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG302 = 1.5 tons, WG101 = 185 tons', style='List Bullet')
doc.add_paragraph('Step 1: pre_cleaning_screening = (1.5 / 185) × 100 = 0.81%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(0.81, 20) = 0.81%', style='List Bullet')
doc.add_paragraph('Result: 0.81%', style='List Bullet')

# KPI 12: 1st Break Capacity per Hour
doc.add_heading('12. 1st Break Capacity per Hour (t/h)', 2)
doc.add_paragraph('Formula: WG202 / net_hours')
doc.add_paragraph('Scales Used: WG202, net_hours')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 310-313\n'
    'if net_hours > 0 and WG202 > 0:\n'
    '    first_break_capacity_per_hour = WG202 / net_hours\n'
    '    # Cap at reasonable maximum (30 t/h)\n'
    '    first_break_capacity_per_hour = min(first_break_capacity_per_hour, 30.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG202 = 200 tons, net_hours = 7.5 hours', style='List Bullet')
doc.add_paragraph('Step 1: first_break_capacity = 200 / 7.5 = 26.67 t/h', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(26.67, 30) = 26.67 t/h', style='List Bullet')
doc.add_paragraph('Result: 26.67 t/h', style='List Bullet')

# KPI 13: Bran Extraction
doc.add_heading('13. Bran Extraction (%)', 2)
doc.add_paragraph('Formula: (WG503 / WG202) × 100')
doc.add_paragraph('Scales Used: WG503, WG202')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 229-230\n'
    'if has_milling_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 317-320\n'
    'if WG202 > 0 and WG503 > 0:\n'
    '    bran_extraction = (WG503 / WG202) * 100.0\n'
    '    # Cap at reasonable maximum (25%)\n'
    '    bran_extraction = min(bran_extraction, 25.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: WG503 = 48 tons, WG202 = 200 tons', style='List Bullet')
doc.add_paragraph('Step 1: bran_extraction = (48 / 200) × 100 = 24%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(24, 25) = 24%', style='List Bullet')
doc.add_paragraph('Result: 24%', style='List Bullet')

doc.add_page_break()

# ============================================================================
# PACKING KPIs
# ============================================================================

doc.add_heading('PACKING KPIs', 1)

# KPI 14: Packing Line Capacity (bags/hr)
doc.add_heading('14. Packing Line Capacity (bags/hr)', 2)
doc.add_paragraph('Formula: PL601 / net_hours')
doc.add_paragraph('Scales Used: PL601_TOT, net_hours')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 333-334\n'
    'if has_packing_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 335-338\n'
    'if net_hours > 0 and PL601 > 0:\n'
    '    packing_capacity = (PL601 / net_hours)\n'
    '    # Cap at reasonable maximum (2000 bags/hr)\n'
    '    packing_capacity = min(packing_capacity, 2000.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: PL601 = 7500 bags, net_hours = 7.5 hours', style='List Bullet')
doc.add_paragraph('Step 1: packing_capacity = 7500 / 7.5 = 1000 bags/hr', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(1000, 2000) = 1000 bags/hr', style='List Bullet')
doc.add_paragraph('Result: 1000 bags/hr', style='List Bullet')

# KPI 15: Daily Packing Output (bags)
doc.add_heading('15. Daily Packing Output (bags)', 2)
doc.add_paragraph('Formula: PL601')
doc.add_paragraph('Scales Used: PL601_TOT')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 333-334\n'
    'if has_packing_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 341-346\n'
    'if PL601 > 0:\n'
    '    daily_packing_output = PL601\n'
    '    # Cap at reasonable maximum (100000 bags per day)\n'
    '    daily_packing_output = min(daily_packing_output, 100000.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: PL601 = 7500 bags', style='List Bullet')
doc.add_paragraph('Step 1: daily_packing_output = 7500 bags', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(7500, 100000) = 7500 bags', style='List Bullet')
doc.add_paragraph('Result: 7500 bags', style='List Bullet')

# KPI 16: Net Hours (Packing)
doc.add_heading('16. Net Hours (hrs)', 2)
doc.add_paragraph('Formula: max(run_hours - downtime, 0.0)')
doc.add_paragraph('Scales Used: run_hours, downtime')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Same as milling net hours'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Line 204\n'
    'net_hours = max(run_hours - downtime, 0.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: run_hours = 8 hours, downtime = 0.5 hours', style='List Bullet')
doc.add_paragraph('Result: 7.5 hours (same calculation as milling)', style='List Bullet')

# KPI 17: Downtime (Packing)
doc.add_heading('17. Downtime (hrs)', 2)
doc.add_paragraph('Formula: Direct SCADA value')
doc.add_paragraph('Scales Used: downtime')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# No validation - direct value from SCADA'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Line 133\n'
    'downtime = safe(row.get("WG202_Stop_Start"), 0.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: downtime = 0.5', style='List Bullet')
doc.add_paragraph('Result: 0.5 hours', style='List Bullet')

# KPI 18: Machine Utilization
doc.add_heading('18. Machine Utilization (%)', 2)
doc.add_paragraph('Formula: (net_hours / daily_hrs) × 100')
doc.add_paragraph('Scales Used: net_hours, daily_hrs')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 333-334\n'
    'if has_packing_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 349-352\n'
    'if daily_hrs > 0 and net_hours > 0:\n'
    '    packing_util = (net_hours / daily_hrs * 100.0)\n'
    '    # Cap at 100% maximum\n'
    '    packing_util = min(packing_util, 100.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: net_hours = 7.5 hours, daily_hrs = 24 hours', style='List Bullet')
doc.add_paragraph('Step 1: packing_util = (7.5 / 24) × 100 = 31.25%', style='List Bullet')
doc.add_paragraph('Step 2: Apply cap - min(31.25, 100) = 31.25%', style='List Bullet')
doc.add_paragraph('Result: 31.25%', style='List Bullet')

# KPI 19: Packing Line Capacity (tons/hr)
doc.add_heading('19. Packing Line Capacity (tons/hr)', 2)
doc.add_paragraph('Formula: (PL601×0.045 + PL602×0.045 + PL603×0.040 + PL606×0.001 + PL607×0.010) / net_hours')
doc.add_paragraph('Scales Used: PL601_TOT, PL602_TOT, PL603_TOT, PL606_TOT, PL607_TOT, net_hours')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 333-334\n'
    'if has_packing_data and has_valid_data:\n'
    '    # Only calculate if validation passes'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 359-374\n'
    'if net_hours > 0:\n'
    '    # Get additional packing data\n'
    '    PL606 = safe(row.get("PL606_TOT", 0.0))\n'
    '    PL607 = safe(row.get("PL607_TOT", 0.0))\n'
    '    \n'
    '    # Convert bags to tons\n'
    '    pl601_tons = PL601 * 0.045  # 45 KG bags\n'
    '    pl602_tons = PL602 * 0.045  # 45 KG bags\n'
    '    pl603_tons = PL603 * 0.040  # 40 KG bran bags\n'
    '    pl606_tons = PL606 * 0.001  # 1 KG bags\n'
    '    pl607_tons = PL607 * 0.010  # 10 KG bags\n'
    '    \n'
    '    total_packing_tons = pl601_tons + pl602_tons + pl603_tons + pl606_tons + pl607_tons\n'
    '    packing_line_capacity_tons_per_hour = total_packing_tons / net_hours\n'
    '    # Cap at reasonable maximum (50 t/h)\n'
    '    packing_line_capacity_tons_per_hour = min(packing_line_capacity_tons_per_hour, 50.0)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input: PL601=5000, PL602=2000, PL603=500, PL606=100, PL607=200, net_hours=7.5h', style='List Bullet')
doc.add_paragraph('Step 1: Convert to tons', style='List Bullet')
doc.add_paragraph('  - PL601: 5000 × 0.045 = 225 tons', style='List Bullet')
doc.add_paragraph('  - PL602: 2000 × 0.045 = 90 tons', style='List Bullet')
doc.add_paragraph('  - PL603: 500 × 0.040 = 20 tons', style='List Bullet')
doc.add_paragraph('  - PL606: 100 × 0.001 = 0.1 tons', style='List Bullet')
doc.add_paragraph('  - PL607: 200 × 0.010 = 2 tons', style='List Bullet')
doc.add_paragraph('Step 2: total_packing_tons = 225 + 90 + 20 + 0.1 + 2 = 337.1 tons', style='List Bullet')
doc.add_paragraph('Step 3: capacity = 337.1 / 7.5 = 44.95 t/h', style='List Bullet')
doc.add_paragraph('Step 4: Apply cap - min(44.95, 50) = 44.95 t/h', style='List Bullet')
doc.add_paragraph('Result: 44.95 t/h', style='List Bullet')

doc.add_page_break()

# ============================================================================
# WATER KPIs
# ============================================================================

doc.add_heading('WATER KPIs', 1)

# KPI 20: Total Pre-Cleaning Water
doc.add_heading('20. Total Pre-Cleaning Water (m³)', 2)
doc.add_paragraph('Formula: DM101 + DM102 (incremental delta)')
doc.add_paragraph('Scales Used: DM101, DM102')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 931-932\n'
    'if current_scada_after_check and baseline:\n'
    '    # Calculate delta (incremental)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 933-954\n'
    'dm101_current = safe(current_scada_after_check.get("DM101", 0.0))\n'
    'dm102_current = safe(current_scada_after_check.get("DM102", 0.0))\n'
    '\n'
    'dm101_baseline = safe(baseline.get("DM101", 0.0))\n'
    'dm102_baseline = safe(baseline.get("DM102", 0.0))\n'
    '\n'
    'dm101_delta = max(0.0, dm101_current - dm101_baseline)\n'
    'dm102_delta = max(0.0, dm102_current - dm102_baseline)\n'
    '\n'
    'totalPreCleaningWater = dm101_delta + dm102_delta'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input (Current): DM101 = 50 m³, DM102 = 30 m³', style='List Bullet')
doc.add_paragraph('Input (Baseline): DM101 = 35 m³, DM102 = 20 m³', style='List Bullet')
doc.add_paragraph('Step 1: dm101_delta = 50 - 35 = 15 m³', style='List Bullet')
doc.add_paragraph('Step 2: dm102_delta = 30 - 20 = 10 m³', style='List Bullet')
doc.add_paragraph('Step 3: totalPreCleaningWater = 15 + 10 = 25 m³', style='List Bullet')
doc.add_paragraph('Result: 25 m³ (incremental)', style='List Bullet')

# KPI 21: Water Clean Wheat
doc.add_heading('21. Water Clean Wheat (m³)', 2)
doc.add_paragraph('Formula: DM201 + DM202 + DM203 (incremental delta)')
doc.add_paragraph('Scales Used: DM201, DM202, DM203')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 931-932\n'
    'if current_scada_after_check and baseline:\n'
    '    # Calculate delta (incremental)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 933-954\n'
    'dm201_current = safe(current_scada_after_check.get("DM201", 0.0))\n'
    'dm202_current = safe(current_scada_after_check.get("DM202", 0.0))\n'
    'dm203_current = safe(current_scada_after_check.get("DM203", 0.0))\n'
    '\n'
    'dm201_baseline = safe(baseline.get("DM201", 0.0))\n'
    'dm202_baseline = safe(baseline.get("DM202", 0.0))\n'
    'dm203_baseline = safe(baseline.get("DM203", 0.0))\n'
    '\n'
    'dm201_delta = max(0.0, dm201_current - dm201_baseline)\n'
    'dm202_delta = max(0.0, dm202_current - dm202_baseline)\n'
    'dm203_delta = max(0.0, dm203_current - dm203_baseline)\n'
    '\n'
    'waterCleanWheat = dm201_delta + dm202_delta + dm203_delta'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Input (Current): DM201 = 40 m³, DM202 = 35 m³, DM203 = 25 m³', style='List Bullet')
doc.add_paragraph('Input (Baseline): DM201 = 20 m³, DM202 = 20 m³, DM203 = 15 m³', style='List Bullet')
doc.add_paragraph('Step 1: dm201_delta = 40 - 20 = 20 m³', style='List Bullet')
doc.add_paragraph('Step 2: dm202_delta = 35 - 20 = 15 m³', style='List Bullet')
doc.add_paragraph('Step 3: dm203_delta = 25 - 15 = 10 m³', style='List Bullet')
doc.add_paragraph('Step 4: waterCleanWheat = 20 + 15 + 10 = 45 m³', style='List Bullet')
doc.add_paragraph('Result: 45 m³ (incremental)', style='List Bullet')

# KPI 22: Total Water Used
doc.add_heading('22. Total Water Used (m³)', 2)
doc.add_paragraph('Formula: DM101 + DM102 + DM201 + DM202 + DM203 (incremental delta)')
doc.add_paragraph('Scales Used: DM101, DM102, DM201, DM202, DM203')

p = doc.add_paragraph()
p.add_run('Validation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 931-932\n'
    'if current_scada_after_check and baseline:\n'
    '    # Calculate delta (incremental)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Calculation Code:').bold = True
code = doc.add_paragraph(
    '# Lines 933-954\n'
    '# Calculate all deltas as shown in previous KPIs\n'
    'dm101_delta = max(0.0, dm101_current - dm101_baseline)\n'
    'dm102_delta = max(0.0, dm102_current - dm102_baseline)\n'
    'dm201_delta = max(0.0, dm201_current - dm201_baseline)\n'
    'dm202_delta = max(0.0, dm202_current - dm202_baseline)\n'
    'dm203_delta = max(0.0, dm203_current - dm203_baseline)\n'
    '\n'
    'totalWaterUsed = (dm101_delta + dm102_delta + \n'
    '                  dm201_delta + dm202_delta + dm203_delta)'
)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Example:').bold = True
doc.add_paragraph('Using deltas from previous examples:', style='List Bullet')
doc.add_paragraph('  - dm101_delta = 15 m³', style='List Bullet')
doc.add_paragraph('  - dm102_delta = 10 m³', style='List Bullet')
doc.add_paragraph('  - dm201_delta = 20 m³', style='List Bullet')
doc.add_paragraph('  - dm202_delta = 15 m³', style='List Bullet')
doc.add_paragraph('  - dm203_delta = 10 m³', style='List Bullet')
doc.add_paragraph('Step 1: totalWaterUsed = 15 + 10 + 20 + 15 + 10 = 70 m³', style='List Bullet')
doc.add_paragraph('Result: 70 m³ (incremental)', style='List Bullet')

# Save document
doc.save('KPI_FORMULAS.docx')
print("✅ KPI_FORMULAS.docx created successfully!")
print("📄 Document contains 22 KPIs with formulas, code, and examples")
